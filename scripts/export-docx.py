#!/usr/bin/env python3
"""Export the finalized resume HTML to an editable, native A4 DOCX.

The HTML is the shared content source. This exporter understands the semantic
classes in template/resume.html and recreates the fixed default layout with
native Word paragraphs, tab stops, borders, numbering, and optional images.

Usage:
    python export-docx.py resume.html output.docx
"""

from __future__ import annotations

import argparse
import re
import sys
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote, urlparse
from zipfile import ZipFile

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.image.image import Image as DocxImage
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - depends on host runtime
    raise SystemExit(
        "缺少 python-docx。请使用当前 Agent 自带的文档运行时执行本脚本，"
        "不要要求最终用户手动安装依赖。"
    ) from exc


IGNORED_TAGS = {"head", "script", "style", "svg"}
VOID_TAGS = {"area", "base", "br", "col", "embed", "hr", "img", "input", "link", "meta", "source", "wbr"}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Default A4 resume style contract. Keep these values synchronized with
# template/resume.html and references/methodology/默认简历版式规范.md.
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_TOP_MM = 9
MARGIN_BOTTOM_MM = 8
MARGIN_LEFT_MM = 7.5
MARGIN_RIGHT_MM = 7.5
# Keep a 2 mm safety allowance around the 40 mm school mark. Word-compatible
# renderers can otherwise clip an image that exactly matches its cell width.
HEADER_COLUMNS_MM = (42, 113, 40)
ENTRY_COLUMNS_MM = (154, 41)
LOGO_BOX_MM = (40, 16)
PORTRAIT_BOX_MM = (20, 27)
LATIN_FONT = "Times New Roman"
EAST_ASIA_FONT = "SimSun"
BODY_SIZE_PT = 10
NAME_SIZE_PT = 18
SECTION_SIZE_PT = 12.5
SECTION_RULE_EIGHTH_POINTS = 10  # 1.25 pt


@dataclass
class Node:
    tag: str
    attrs: dict[str, str] = field(default_factory=dict)
    children: list["Node | str"] = field(default_factory=list)

    @property
    def classes(self) -> set[str]:
        return set(self.attrs.get("class", "").split())


class DOMParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = Node("document")
        self.stack = [self.root]

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        node = Node(tag.lower(), {key: value or "" for key, value in attrs})
        self.stack[-1].children.append(node)
        if tag.lower() not in VOID_TAGS:
            self.stack.append(node)

    def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self.handle_starttag(tag, attrs)
        if tag.lower() not in VOID_TAGS:
            self.handle_endtag(tag)

    def handle_endtag(self, tag: str) -> None:
        for index in range(len(self.stack) - 1, 0, -1):
            if self.stack[index].tag == tag.lower():
                del self.stack[index:]
                return

    def handle_data(self, data: str) -> None:
        self.stack[-1].children.append(data)


def walk(node: Node) -> Iterable[Node]:
    for child in node.children:
        if isinstance(child, Node):
            yield child
            if child.tag not in IGNORED_TAGS:
                yield from walk(child)


def find_class(node: Node, class_name: str) -> Node | None:
    return next((candidate for candidate in walk(node) if class_name in candidate.classes), None)


def find_tag(node: Node, tag: str) -> Node | None:
    return next((candidate for candidate in walk(node) if candidate.tag == tag), None)


def find_all(node: Node | None, *, tag: str | None = None, class_name: str | None = None) -> list[Node]:
    if node is None:
        return []
    return [
        candidate
        for candidate in walk(node)
        if (tag is None or candidate.tag == tag)
        and (class_name is None or class_name in candidate.classes)
    ]


def direct_children(node: Node, *, tag: str | None = None, class_name: str | None = None) -> list[Node]:
    return [
        child
        for child in node.children
        if isinstance(child, Node)
        and (tag is None or child.tag == tag)
        and (class_name is None or class_name in child.classes)
    ]


def text_of(node: Node | None) -> str:
    if node is None:
        return ""
    chunks: list[str] = []

    def collect(current: Node) -> None:
        if current.tag in IGNORED_TAGS:
            return
        for child in current.children:
            if isinstance(child, str):
                chunks.append(child)
            elif child.tag == "br":
                chunks.append("\n")
            else:
                collect(child)

    collect(node)
    return re.sub(r"[\t\r\f\v ]+", " ", "".join(chunks).replace("\xa0", " ")).strip()


def inline_segments(node: Node) -> list[tuple[str, bool]]:
    chars: list[tuple[str, bool]] = []

    def collect(current: Node, bold: bool = False) -> None:
        if current.tag in IGNORED_TAGS:
            return
        is_bold = bold or current.tag in {"b", "strong"}
        for child in current.children:
            if isinstance(child, str):
                chars.extend((character, is_bold) for character in child.replace("\xa0", " "))
            elif child.tag == "br":
                chars.append(("\n", is_bold))
            else:
                collect(child, is_bold)

    collect(node)
    normalized: list[tuple[str, bool]] = []
    pending_space = False
    for character, bold in chars:
        if character.isspace() and character != "\n":
            pending_space = bool(normalized)
            continue
        if pending_space and character != "\n":
            normalized.append((" ", bold))
        pending_space = False
        normalized.append((character, bold))
    while normalized and normalized[-1][0].isspace():
        normalized.pop()

    segments: list[tuple[str, bool]] = []
    for character, bold in normalized:
        if segments and segments[-1][1] == bold:
            segments[-1] = (segments[-1][0] + character, bold)
        else:
            segments.append((character, bold))
    return segments


def set_style_fonts(style, latin: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT) -> None:
    style.font.name = latin
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key, value in (("ascii", latin), ("hAnsi", latin), ("eastAsia", east_asia), ("cs", latin)):
        r_fonts.set(qn(f"w:{key}"), value)


def set_run_fonts(run, latin: str = LATIN_FONT, east_asia: str = EAST_ASIA_FONT) -> None:
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key, value in (("ascii", latin), ("hAnsi", latin), ("eastAsia", east_asia), ("cs", latin)):
        r_fonts.set(qn(f"w:{key}"), value)


def add_segments(paragraph, segments: list[tuple[str, bool]]) -> None:
    for content, bold in segments:
        run = paragraph.add_run(content)
        run.bold = bold
        set_run_fonts(run)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(SECTION_RULE_EIGHTH_POINTS))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "111111")
    borders.append(bottom)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def set_table_widths(table, widths_mm: tuple[float, ...]) -> None:
    if len(table.columns) != len(widths_mm):
        raise ValueError("表格列数与宽度数量不一致。")
    widths = [round(width / 25.4 * 1440) for width in widths_mm]
    table.autofit = False
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_style_fonts(normal)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.02

    def paragraph_style(name: str, *, size: float, bold: bool = False):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        set_style_fonts(style)
        style.font.size = Pt(size)
        style.font.bold = bold
        return style

    name_style = paragraph_style("Resume Name", size=NAME_SIZE_PT, bold=True)
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_style.paragraph_format.space_after = Pt(5)
    name_style.paragraph_format.keep_with_next = True

    contact_style = paragraph_style("Resume Contact", size=BODY_SIZE_PT)
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_style.paragraph_format.keep_with_next = True

    availability_style = paragraph_style("Resume Availability", size=BODY_SIZE_PT)
    availability_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    availability_style.paragraph_format.space_before = Pt(1)
    availability_style.paragraph_format.keep_with_next = True

    section_style = paragraph_style("Resume Section", size=SECTION_SIZE_PT, bold=True)
    section_style.paragraph_format.space_before = Pt(5.5)
    section_style.paragraph_format.space_after = Pt(3)
    section_style.paragraph_format.keep_with_next = True

    entry_style = paragraph_style("Resume Entry Header", size=BODY_SIZE_PT, bold=False)
    entry_style.paragraph_format.space_after = Pt(1)
    entry_style.paragraph_format.keep_with_next = True

    detail_style = paragraph_style("Resume Detail", size=BODY_SIZE_PT)
    detail_style.paragraph_format.space_after = Pt(0.8)
    detail_style.paragraph_format.line_spacing = 1.0
    detail_style.paragraph_format.keep_together = True

    bullet = styles["List Bullet"]
    set_style_fonts(bullet)
    bullet.font.size = Pt(BODY_SIZE_PT)
    bullet.paragraph_format.left_indent = Mm(4.8)
    bullet.paragraph_format.first_line_indent = Mm(-3.2)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(0.8)
    bullet.paragraph_format.line_spacing = 1.0
    bullet.paragraph_format.keep_together = True


def resolve_local_image(node: Node | None, html_path: Path) -> Path | None:
    if node is None or not node.attrs.get("src"):
        return None
    raw = node.attrs["src"].strip()
    parsed = urlparse(raw)
    if parsed.scheme not in {"", "file"}:
        raise ValueError("DOCX 只接受本地校徽/照片文件，不能直接下载远程图片。")
    path = Path(unquote(parsed.path)) if parsed.scheme == "file" else Path(unquote(raw))
    if not path.is_absolute():
        path = html_path.parent / path
    path = path.resolve()
    if not path.is_file():
        raise ValueError(f"图片不存在: {path}")
    if path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
        raise ValueError(f"DOCX 图片仅支持 PNG/JPG: {path.name}")
    return path


def image_ratio(path: Path) -> float:
    image = DocxImage.from_file(str(path))
    if image.width <= 0 or image.height <= 0:
        raise ValueError(f"无法读取图片尺寸: {path}")
    return image.width / image.height


def assert_standard_portrait(path: Path) -> None:
    actual_ratio = image_ratio(path)
    width_mm, height_mm = PORTRAIT_BOX_MM
    target_ratio = width_mm / height_mm
    if abs(actual_ratio / target_ratio - 1) > 0.01:
        raise ValueError(
            f"证件照尚未转换为 {width_mm:g} x {height_mm:g} mm 固定画布。"
            "请先运行 scripts/prepare-resume-image.py portrait。"
        )


def fit_image_box(path: Path, maximum_width_mm: float, maximum_height_mm: float) -> tuple[float, float]:
    ratio = image_ratio(path)
    if ratio >= maximum_width_mm / maximum_height_mm:
        return maximum_width_mm, maximum_width_mm / ratio
    return maximum_height_mm * ratio, maximum_height_mm


def add_header(document, page: Node, html_path: Path, expected: list[str]) -> None:
    name = text_of(find_class(page, "name"))
    if not name:
        raise ValueError("未识别到 .name；请使用 template/resume.html 的固定结构。")
    expected.append(name)

    table = document.add_table(rows=1, cols=3)
    set_table_widths(table, HEADER_COLUMNS_MM)
    remove_table_borders(table)
    table.rows[0].height = Mm(28)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)

    logo = resolve_local_image(find_class(page, "school-logo"), html_path)
    portrait = resolve_local_image(find_class(page, "portrait"), html_path)
    if logo:
        logo_width, logo_height = fit_image_box(logo, *LOGO_BOX_MM)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run().add_picture(
            str(logo), width=Mm(logo_width), height=Mm(logo_height)
        )
    else:
        # Keep an explicit paragraph payload so Word-compatible renderers do
        # not collapse the optional left column.
        table.cell(0, 0).paragraphs[0].add_run("\u00a0")
    if portrait:
        assert_standard_portrait(portrait)
        paragraph = table.cell(0, 2).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run().add_picture(
            str(portrait), width=Mm(PORTRAIT_BOX_MM[0]), height=Mm(PORTRAIT_BOX_MM[1])
        )
    else:
        # LibreOffice may suppress the center cell when a one-sided header
        # table ends in a structurally empty cell.
        table.cell(0, 2).paragraphs[0].add_run("\u00a0")

    center = table.cell(0, 1)
    name_paragraph = center.paragraphs[0]
    name_paragraph.style = document.styles["Resume Name"]
    name_run = name_paragraph.add_run(name)
    name_run.bold = True
    set_run_fonts(name_run)

    contact = find_class(page, "contact")
    contact_items = [text_of(span) for span in find_all(contact, tag="span")]
    contact_items = [item for item in contact_items if item]
    if contact_items:
        expected.extend(contact_items)
        paragraph = center.add_paragraph(style="Resume Contact")
        run = paragraph.add_run("｜".join(contact_items))
        set_run_fonts(run)

    availability = text_of(find_class(page, "availability"))
    if availability:
        expected.append(availability)
        paragraph = center.add_paragraph(style="Resume Availability")
        run = paragraph.add_run(availability)
        set_run_fonts(run)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.5


def build_docx(html_path: Path, output_path: Path) -> tuple[int, int]:
    parser = DOMParser()
    parser.feed(html_path.read_text(encoding="utf-8"))
    page = find_class(parser.root, "page") or parser.root
    sections = find_all(page, tag="section")
    if not sections:
        raise ValueError("未识别到 section；请使用 template/resume.html 的固定结构。")

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)
    configure_styles(document)

    expected: list[str] = []
    add_header(document, page, html_path, expected)
    bullet_count = 0

    for source_section in sections:
        heading = text_of(find_tag(source_section, "h2"))
        if not heading:
            continue
        expected.append(heading)
        paragraph = document.add_paragraph(style="Resume Section")
        run = paragraph.add_run(heading)
        run.bold = True
        set_run_fonts(run)
        add_bottom_border(paragraph)

        for entry in direct_children(source_section, class_name="entry"):
            header = find_class(entry, "entry-head")
            if header:
                title = text_of(find_class(header, "entry-title"))
                role = text_of(find_class(header, "entry-role"))
                date = text_of(find_class(header, "entry-date"))
                expected.extend(value for value in (title, role, date) if value)
                header_table = document.add_table(rows=1, cols=2)
                set_table_widths(header_table, ENTRY_COLUMNS_MM)
                remove_table_borders(header_table)
                left_cell, right_cell = header_table.rows[0].cells
                for cell in (left_cell, right_cell):
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                paragraph = left_cell.paragraphs[0]
                paragraph.style = document.styles["Resume Entry Header"]
                if title:
                    run = paragraph.add_run(title)
                    run.bold = True
                    set_run_fonts(run)
                if role:
                    run = paragraph.add_run(f" ｜ {role}")
                    run.bold = True
                    set_run_fonts(run)
                if date:
                    date_paragraph = right_cell.paragraphs[0]
                    date_paragraph.style = document.styles["Resume Entry Header"]
                    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = date_paragraph.add_run(date)
                    run.font.size = Pt(BODY_SIZE_PT)
                    set_run_fonts(run)

            for detail in direct_children(entry, class_name="kv"):
                expected.append(text_of(detail))
                paragraph = document.add_paragraph(style="Resume Detail")
                add_segments(paragraph, inline_segments(detail))

            for item in find_all(entry, tag="li"):
                expected.append(text_of(item))
                paragraph = document.add_paragraph(style="List Bullet")
                add_segments(paragraph, inline_segments(item))
                bullet_count += 1

        for detail in direct_children(source_section, class_name="kv"):
            expected.append(text_of(detail))
            paragraph = document.add_paragraph(style="Resume Detail")
            add_segments(paragraph, inline_segments(detail))

    properties = document.core_properties
    properties.title = f"{expected[0]} - 简历"
    properties.subject = ""
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.keywords = ""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    validate_docx(output_path, expected, bullet_count)
    return len(sections), bullet_count


def validate_docx(output_path: Path, expected: list[str], bullet_count: int) -> None:
    with ZipFile(output_path) as archive:
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        missing = required.difference(archive.namelist())
        if missing:
            raise ValueError(f"DOCX 包缺少必要文件: {', '.join(sorted(missing))}")
        root = ET.fromstring(archive.read("word/document.xml"))
        paragraphs = [
            "".join(text.text or "" for text in paragraph.findall(".//w:t", {"w": W_NS}))
            for paragraph in root.findall(".//w:p", {"w": W_NS})
        ]
        all_text = re.sub(r"\s+", " ", "\n".join(paragraphs)).strip()
        missing_text = [
            fragment for fragment in expected
            if fragment and re.sub(r"\s+", " ", fragment).strip() not in all_text
        ]
        if missing_text:
            raise ValueError(f"DOCX 内容校验失败，缺少: {'；'.join(missing_text[:3])}")

        page_size = root.find(".//w:pgSz", {"w": W_NS})
        if page_size is None:
            raise ValueError("DOCX 结构校验失败：没有页面尺寸。")
        width = int(page_size.get(qn("w:w"), "0"))
        height = int(page_size.get(qn("w:h"), "0"))
        if abs(width - 11906) > 2 or abs(height - 16838) > 2:
            raise ValueError(f"DOCX 结构校验失败：页面不是 A4（{width} x {height} twips）。")

        margins = root.find(".//w:pgMar", {"w": W_NS})
        if margins is None:
            raise ValueError("DOCX 结构校验失败：没有页边距。")
        expected_margins = {
            "top": round(MARGIN_TOP_MM / 25.4 * 1440),
            "bottom": round(MARGIN_BOTTOM_MM / 25.4 * 1440),
            "left": round(MARGIN_LEFT_MM / 25.4 * 1440),
            "right": round(MARGIN_RIGHT_MM / 25.4 * 1440),
        }
        for side, expected_value in expected_margins.items():
            actual_value = int(margins.get(qn(f"w:{side}"), "0"))
            if abs(actual_value - expected_value) > 2:
                raise ValueError(
                    f"DOCX 结构校验失败：{side} 页边距为 {actual_value} twips，"
                    f"应为 {expected_value} twips。"
                )

        if bullet_count and "word/numbering.xml" not in archive.namelist():
            raise ValueError("DOCX 结构校验失败：项目符号不是 Word 原生编号。")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将定稿简历 HTML 导出为固定版式的可编辑 DOCX。")
    parser.add_argument("html", type=Path, help="输入 resume.html")
    parser.add_argument("output", type=Path, nargs="?", help="输出 .docx；默认与 HTML 同名")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    html_path = args.html.expanduser().resolve()
    output_path = (args.output or html_path.with_suffix(".docx")).expanduser().resolve()
    if not html_path.is_file():
        print(f"错误: HTML 不存在: {html_path}", file=sys.stderr)
        return 2
    if output_path.suffix.lower() != ".docx":
        print("错误: 输出文件必须使用 .docx 扩展名。", file=sys.stderr)
        return 2
    try:
        section_count, bullet_count = build_docx(html_path, output_path)
    except (OSError, ValueError) as exc:
        print(f"错误: {exc}", file=sys.stderr)
        return 3
    print(f"DOCX: {output_path}")
    print(f"结构: {section_count} 个区块，{bullet_count} 条项目符号")
    print("下一步: 将 DOCX 渲染为逐页 PNG，确认恰好一页且版式无缺陷。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
